"""
pipeline/documents.py — Stage: document upload text extraction.

Turns an uploaded PDF / Word / CSV / plain-text / Markdown file into plain
text so it can be handed to the chat pipeline like any other prompt. Kept
deliberately simple: no OCR fallback for scanned/image-only PDFs (that's
what the separate /api/ocr endpoint on an image upload is for), no chunking
or embeddings-based retrieval yet (that's the ChromaDB step on the roadmap)
— this just gets a document's raw text in front of the model in one shot,
capped to a safe length.
"""

import csv
import io

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt", "md", "markdown", "csv"}


def extract_text(content: bytes, filename: str) -> str:
    ext = (filename.rsplit(".", 1)[-1].lower() if "." in filename else "")
    try:
        if ext == "pdf":
            return _extract_pdf(content)
        if ext == "docx":
            return _extract_docx(content)
        if ext == "csv":
            return _extract_csv(content)
        if ext in ("txt", "md", "markdown"):
            return content.decode("utf-8", errors="replace")
        # Unknown extension — best-effort decode as text rather than failing outright.
        return content.decode("utf-8", errors="replace")
    except Exception as e:
        return f"[Could not extract text from '{filename}': {e}]"


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    parts = []
    for page in reader.pages[:40]:  # cap very long PDFs
        parts.append(page.extract_text() or "")
    text = "\n\n".join(parts).strip()
    if not text:
        return "[This PDF has no extractable text — it may be a scanned/image-only document. Try the Extract Text (OCR) option instead by uploading a page as an image.]"
    return text


def _extract_docx(content: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def _extract_csv(content: bytes) -> str:
    text = content.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    preview = rows[:200]  # cap very large spreadsheets
    out = "\n".join(",".join(cell for cell in row) for row in preview)
    if len(rows) > 200:
        out += f"\n... ({len(rows) - 200} more rows not shown, {len(rows)} total)"
    return out
